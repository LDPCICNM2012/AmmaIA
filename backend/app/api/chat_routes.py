import io
from fastapi import APIRouter, HTTPException, Depends, UploadFile, File, Form
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from typing import Optional, List, Dict, Any
import json
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

from ..auth.security import get_current_user
from ..database.db import (
    verificar_y_consumir_cuota,
    obtener_estado_cuota,
    guardar_chat_db,
    obtener_chats_usuario,
    borrar_chat_db
)
from ..rag.rag_engine import consultar_ammaia
from ..rag.doc_parser import extraer_texto_de_archivo
from ..rag.pdf_generator import generar_dictamen_pdf
from ..config import DEFAULT_MODEL

router = APIRouter(prefix="/chat", tags=["Consultas Jurídicas RAG"])

class ChatConsultaRequest(BaseModel):
    pregunta: str
    historial: Optional[List[Dict[str, Any]]] = []
    modelo: Optional[str] = DEFAULT_MODEL
    chat_id: Optional[str] = None
    texto_adjunto: Optional[str] = None
    nombre_adjunto: Optional[str] = None

class GuardarChatRequest(BaseModel):
    chat_id: str
    titulo: str
    mensajes: List[Dict[str, Any]]

class ExportarDocRequest(BaseModel):
    titulo: str
    mensajes: List[Dict[str, Any]]

@router.post("/consultar")
def endpoint_consultar_ammaia(req: ChatConsultaRequest, user: Dict[str, Any] = Depends(get_current_user)):
    user_id = user["id"]
    is_premium = user["is_premium"]

    # 1. Verificar y consumir cuota (5 al día para gratis, ilimitado para Premium)
    puede_consultar, usados, restantes = verificar_y_consumir_cuota(user_id, is_premium)
    if not puede_consultar:
        raise HTTPException(
            status_code=429,
            detail="Has alcanzado el límite de 5 consultas gratuitas por hoy. Vuelve mañana a las 00:00 o actualiza al Plan Premium Jurídico para consultas ilimitadas."
        )

    # Si hay texto de archivo adjunto (PDF, Word, etc.), anexarlo a la pregunta
    pregunta_con_contexto = req.pregunta
    if req.texto_adjunto:
        nombre_doc = req.nombre_adjunto or "Documento Adjunto"
        pregunta_con_contexto = (
            f"=== DOCUMENTO ADJUNTO POR EL LETRADO ({nombre_doc}) ===\n"
            f"{req.texto_adjunto}\n"
            f"=== FIN DOCUMENTO ADJUNTO ===\n\n"
            f"CONSULTA ESPECÍFICA SOBRE EL DOCUMENTO O CASO:\n{req.pregunta}"
        )

    # 2. Ejecutar motor RAG con Gemini
    resultado_rag = consultar_ammaia(
        pregunta=pregunta_con_contexto,
        historial=req.historial,
        modelo=req.modelo or DEFAULT_MODEL
    )

    # 3. Obtener estado actualizado de cuota
    estado_cuota = obtener_estado_cuota(user_id, is_premium)

    return {
        "success": True,
        "respuesta": resultado_rag["respuesta"],
        "fuentes": resultado_rag["fuentes"],
        "modelo": resultado_rag["modelo"],
        "cuota": estado_cuota
    }

@router.post("/subir-archivo")
async def endpoint_subir_archivo(archivo: UploadFile = File(...), user: Dict[str, Any] = Depends(get_current_user)):
    """Extrae el contenido textual de un archivo PDF, DOCX o TXT para adjuntarlo al chat."""
    try:
        contenido_bytes = await archivo.read()
        texto_extraido = extraer_texto_de_archivo(contenido_bytes, archivo.filename)
        return {
            "success": True,
            "nombre_archivo": archivo.filename,
            "caracteres": len(texto_extraido),
            "texto": texto_extraido
        }
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"No se pudo procesar el archivo: {e}")

@router.get("/historial")
def endpoint_obtener_historial(user: Dict[str, Any] = Depends(get_current_user)):
    chats = obtener_chats_usuario(user["id"])
    return {"chats": chats}

@router.post("/guardar")
def endpoint_guardar_chat(req: GuardarChatRequest, user: Dict[str, Any] = Depends(get_current_user)):
    guardar_chat_db(req.chat_id, user["id"], req.titulo, req.mensajes)
    return {"success": True, "mensaje": "Chat guardado correctamente."}

@router.delete("/{chat_id}")
def endpoint_borrar_chat(chat_id: str, user: Dict[str, Any] = Depends(get_current_user)):
    borrar_chat_db(chat_id, user["id"])
    return {"success": True, "mensaje": "Chat eliminado."}

@router.post("/exportar-word")
def endpoint_exportar_word(req: ExportarDocRequest, user: Dict[str, Any] = Depends(get_current_user)):
    """Genera y descarga un documento Word (.docx) formal con el dictamen de AmmaIA."""
    doc = Document()

    # Título
    titulo_p = doc.add_heading(f"Dictamen Jurídico — {req.titulo}", 0)
    titulo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta_p = doc.add_paragraph(f"Generado por AmmaIA • Letrado/a: {user['nombre']} • Fecha: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # Conversación / Dictamen
    for msg in req.mensajes:
        rol = "👤 Consulta del Letrado" if msg.get("role") in ["user", "usuario"] else "⚖️ Dictamen AmmaIA"
        p = doc.add_paragraph()
        r = p.add_run(f"{rol}:\n")
        r.bold = True
        r.font.size = Pt(11)
        if msg.get("role") in ["user", "usuario"]:
            r.font.color.rgb = RGBColor(37, 99, 235)
        else:
            r.font.color.rgb = RGBColor(217, 119, 6)

        doc.add_paragraph(str(msg.get("content", "")))
        doc.add_paragraph("─" * 40)

    # Guardar en memoria
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    filename = f"Dictamen_AmmaIA_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    return StreamingResponse(
        file_stream,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )

@router.post("/exportar-pdf")
def endpoint_exportar_pdf(req: ExportarDocRequest, user: Dict[str, Any] = Depends(get_current_user)):
    """Genera y descarga un archivo PDF formal y maquetado con ReportLab."""
    try:
        pdf_buffer = generar_dictamen_pdf(
            titulo=req.titulo or "Consulta Legal",
            mensajes=req.mensajes,
            usuario_nombre=user["nombre"]
        )
        filename = f"Dictamen_AmmaIA_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        return StreamingResponse(
            pdf_buffer,
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generando PDF: {e}")
